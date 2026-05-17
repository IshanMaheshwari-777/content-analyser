"""
test_app.py — Integration test for DegreeBaba parser, mapper, and validator.
"""

import os
from docx import Document
from parser.extractor import extract_sections, detect_page_type
from parser.heading_mapper import map_headings, get_field_schema
from validator.payload_validator import validate_payload

def run_test():
    print("=== Creating Mock Word Document ===")
    doc = Document()
    
    # Add a heading
    h1 = doc.add_heading("About Amity Online", level=1)
    p1 = doc.add_paragraph("Amity University Online is a leading EdTech university in India offering high quality degrees.")
    
    # Add a second heading (Quick Facts)
    doc.add_heading("University Facts", level=1)
    p2 = doc.add_paragraph("Established: 2005")
    p3 = doc.add_paragraph("NAAC Grade: A++")
    p4 = doc.add_paragraph("Mode of learning: Online")
    p5 = doc.add_paragraph("UGC Status: Approved")
    
    # Add a third heading (FAQs)
    doc.add_heading("Frequently Asked Questions", level=1)
    p6 = doc.add_paragraph("Question: What is the fee structure?")
    p7 = doc.add_paragraph("Answer: The fee structure varies by course, starting from 1.5L.")
    
    # Save document to bytes
    import io
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    print("=== Testing Extractor ===")
    extracted = extract_sections(docx_bytes.read())
    print("Extracted Title:", extracted["document_title"])
    print("Extracted Sections:", [s["heading"] for s in extracted["sections"]])
    assert len(extracted["sections"]) == 3
    assert extracted["document_title"] == "About Amity Online"
    
    print("=== Testing Page Type Detection ===")
    ptype = detect_page_type("Amity_Online_University_Review.docx", extracted["sections"])
    print("Detected Page Type:", ptype)
    assert ptype == "university"
    
    print("=== Testing Heading Mapping & Quick Facts ===")
    mapping_result = map_headings(extracted["sections"], ptype)
    print("Mappings:")
    for m in mapping_result["mappings"]:
        print(f"  {m['heading']} -> {m['acf_field']} ({m['confidence']}% via {m['method']})")
    print("Extracted Quick Facts:", mapping_result["quick_facts"])
    
    # Assertions for mapping
    facts = mapping_result["quick_facts"]
    assert facts.get("established_year") == "2005"
    assert facts.get("naac_grade") == "A++"
    assert facts.get("mode_of_learning") == "Online"
    assert facts.get("ugc_status") == "Approv"
    
    print("=== Testing Payload Validator ===")
    # Create a mock payload
    payload = {
        "university_name": "Amity Online",
        "about_content": "<p>Amity University Online is a leading EdTech university in India offering high quality degrees.</p>",
        "faqs": [{"question": "What is the fee structure?", "answer": "The fee structure varies by course, starting from 1.5L."}],
        **facts
    }
    schema = get_field_schema(ptype)
    validation = validate_payload(payload, ptype, mapping_result, schema)
    print("Validation Score:", validation["score"])
    print("Missing fields:", validation["missing"])
    print("Thin fields:", validation["thin"])
    
    print("=== Integration Test Passed Successfully! ===")

if __name__ == "__main__":
    run_test()
