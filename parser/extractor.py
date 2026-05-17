"""
extractor.py — Word file parsing module for DegreeBaba.

Uses python-docx to walk the document body and split into sections by heading.
Also uses mammoth to convert the full file to HTML and slice by heading boundaries
for clean WYSIWYG content.
"""

import re
import io

from docx import Document
import mammoth


def extract_sections(docx_bytes: bytes) -> dict:
    """
    Parse a .docx file and return structured sections + full HTML.
    
    Args:
        docx_bytes: Raw bytes of the uploaded .docx file.
    
    Returns:
        dict with keys:
            - sections: list of {heading, level, paragraphs, tables, raw_text}
            - full_html: full mammoth HTML output
            - html_sections: dict mapping heading -> HTML slice
            - document_title: extracted title from first heading or bold line
    """
    doc = Document(io.BytesIO(docx_bytes))
    
    # --- Extract sections via python-docx ---
    sections = []
    current_section = None
    document_title = None
    
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        
        if tag == "p":
            para = _find_paragraph(doc, element)
            if para is None:
                continue
                
            # Check if it's a heading
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = _get_heading_level(para.style.name)
                heading_text = para.text.strip()
                
                if not heading_text:
                    continue
                
                # Capture document title from first H1 or first heading
                if document_title is None and level <= 2:
                    document_title = heading_text
                
                current_section = {
                    "heading": heading_text,
                    "level": level,
                    "paragraphs": [],
                    "tables": [],
                    "raw_text": ""
                }
                sections.append(current_section)
            else:
                text = para.text.strip()
                if text:
                    # If no section yet, check for bold text as potential title
                    if current_section is None:
                        if document_title is None and _is_bold(para):
                            document_title = text
                    else:
                        current_section["paragraphs"].append(text)
                        current_section["raw_text"] += text + "\n"
        
        elif tag == "tbl":
            table_data = _extract_table(doc, element)
            if current_section is not None and table_data:
                current_section["tables"].append(table_data)
                # Also add table text to raw_text for AI processing
                for row in table_data:
                    current_section["raw_text"] += " | ".join(str(c) for c in row) + "\n"
    
    # --- Extract full HTML via mammoth ---
    mammoth_result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
    full_html = mammoth_result.value
    
    # --- Slice HTML by heading boundaries ---
    html_sections = _slice_html_by_headings(full_html, sections)
    
    # --- Fallback for document title ---
    if document_title is None and sections:
        document_title = sections[0]["heading"]
    if document_title is None:
        document_title = "Untitled Document"
    
    return {
        "sections": sections,
        "full_html": full_html,
        "html_sections": html_sections,
        "document_title": document_title
    }


def detect_page_type(filename: str, sections: list) -> str:
    """
    Auto-detect page type from filename keywords and heading patterns.
    
    Returns: 'university', 'course', or 'specialization'
    """
    filename_lower = filename.lower()
    headings_lower = [s["heading"].lower() for s in sections]
    all_headings_text = " ".join(headings_lower)
    
    # Specialization signals (check first — most specific)
    spec_signals = [
        "specialization highlights", "other specializations", "capstone",
        "explore other", "course details"
    ]
    spec_score = sum(1 for sig in spec_signals if sig in all_headings_text)
    if "specialization" in filename_lower:
        spec_score += 3
    
    # Course signals
    course_signals = [
        "syllabus", "specializations offered", "who can apply",
        "program highlights", "about the program"
    ]
    course_score = sum(1 for sig in course_signals if sig in all_headings_text)
    if "course" in filename_lower or "program" in filename_lower:
        course_score += 3
    if "mba" in filename_lower or "bba" in filename_lower or "mca" in filename_lower or "bca" in filename_lower:
        course_score += 2
    
    # University signals
    uni_signals = [
        "faculty members", "programs offered", "about the university",
        "university facts", "university detail"
    ]
    uni_score = sum(1 for sig in uni_signals if sig in all_headings_text)
    if "university" in filename_lower or "college" in filename_lower:
        uni_score += 3
    
    scores = {
        "specialization": spec_score,
        "course": course_score,
        "university": uni_score
    }
    
    best = max(scores, key=scores.get)
    if scores[best] == 0:
        # Default to course if no clear signal
        return "course"
    return best


def _find_paragraph(doc, element):
    """Find the python-docx Paragraph object matching this XML element."""
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _get_heading_level(style_name: str) -> int:
    """Extract heading level from style name like 'Heading 1', 'Heading 2'."""
    match = re.search(r"(\d+)", style_name)
    if match:
        return int(match.group(1))
    return 1


def _is_bold(para) -> bool:
    """Check if the first run in a paragraph is bold."""
    if para.runs:
        return para.runs[0].bold is True
    return False


def _extract_table(doc, tbl_element):
    """Extract table data from a table XML element."""
    for table in doc.tables:
        if table._element is tbl_element:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            return rows
    return []


def _slice_html_by_headings(full_html: str, sections: list) -> dict:
    """
    Slice the mammoth HTML output by heading boundaries.
    Maps each heading to its corresponding HTML slice.
    """
    html_sections = {}
    
    if not sections:
        return html_sections
    
    # Find positions of each heading in the HTML sequentially to handle duplicates correctly
    positions = []
    current_search_index = 0
    for section in sections:
        heading_text = section["heading"]
        normalized_heading_text = re.escape(heading_text)
        
        # 1. Try standard h1-h6 tag match
        pattern = re.compile(rf"<h[1-6][^>]*>\s*{normalized_heading_text}\s*</h[1-6]>", re.IGNORECASE)
        match = pattern.search(full_html, current_search_index)
        
        if match:
            positions.append((heading_text, match.start(), match.end()))
            current_search_index = match.end()
        else:
            # 2. Try fallback bold paragraph match (mammoth sometimes parses headings as bold p)
            pattern_fallback = re.compile(rf"<p[^>]*>\s*(?:<strong>|<b>)?\s*{normalized_heading_text}\s*(?:</strong>|</b>)?\s*</p>", re.IGNORECASE)
            match_fb = pattern_fallback.search(full_html, current_search_index)
            if match_fb:
                positions.append((heading_text, match_fb.start(), match_fb.end()))
                current_search_index = match_fb.end()
    
    # Sort by position in the HTML (should already be sorted but safe-guarding)
    positions.sort(key=lambda x: x[1])
    
    # Extract HTML slices between headings
    for i, (heading, start, tag_end) in enumerate(positions):
        if i + 1 < len(positions):
            end = positions[i + 1][1]
        else:
            end = len(full_html)
        
        # Include everything after the heading tag until next heading
        html_slice = full_html[tag_end:end].strip()
        html_sections[heading] = html_slice
    
    return html_sections

